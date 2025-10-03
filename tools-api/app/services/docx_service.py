from typing import List
from io import BytesIO
from docx import Document


def parse_docx_to_text(file_bytes: bytes) -> str:
    """Parse a .docx file (bytes) and return plain text (joined paragraphs).

    Returns a single string with newlines between paragraphs.
    """
    bio = BytesIO(file_bytes)
    doc = Document(bio)
    parts: List[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    return "\n".join(parts)


def create_docx_from_text(text: str) -> bytes:
    """Create a .docx file from plain text. Paragraphs separated by double-newline or single newline preserved."""
    doc = Document()
    # Split into paragraphs by blank lines
    paras = [p for p in text.split("\n\n")]
    for p in paras:
        # preserve internal newlines inside a paragraph by replacing them with line breaks
        # python-docx doesn't support line breaks easily; add as separate paragraphs
        lines = p.split("\n")
        for i, line in enumerate(lines):
            if i == 0:
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
